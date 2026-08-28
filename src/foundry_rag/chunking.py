"""Belgeleri okuma ve chunk'lama.

Kural: paragraf sınırlarında böl. Bir paragraf tek başına üst sınırı aşarsa
cümlelere, cümle de aşarsa karaktere düşülür — hiçbir metin sessizce kaybolmaz.
"""

from __future__ import annotations

import re
from pathlib import Path

from . import config

SUPPORTED_SUFFIXES = {".md", ".txt", ".docx", ".pdf"}

_PARAGRAPH_SPLIT = re.compile(r"\n\s*\n")
_SENTENCE_SPLIT = re.compile(r"(?<=[.!?…])\s+")

# PDF'te gömülü font subset'i çözülemediğinde pypdf karakterleri "/gid00047"
# gibi glif numaralarına çevirir. Bu metin okunamaz çöptür: retrieval'a girip
# bağlama düşerse model onu yorumlamaya çalışıp saçma cevap üretir (ölçüldü).
_GLIF_KODU = re.compile(r"/gid\d{3,}")
# Tek tük geçen bir kod yanlış alarm olabilir; yoğunlaşma aranır.
_GLIF_ESIGI = 3


def read_text(path: Path) -> str:
    """Desteklenen bir dosyayı düz metne çevir."""
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    raise ValueError(f"Desteklenmeyen dosya türü: {path.suffix} ({path.name})")


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - kurulum hatası
        raise RuntimeError(
            "docx okumak için 'python-docx' gerekli: pip install -r requirements.txt"
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(part.strip() for part in parts if part.strip())


def _read_pdf(path: Path) -> str:
    """PDF'ten metin çıkar.

    DİKKAT: yalnızca metin katmanı okunur. Taranmış/fotoğraflanmış PDF'te metin
    katmanı yoktur ve sonuç boş döner (ingest böyle dosyaları atlayıp raporlar) —
    OCR yapılmaz.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - kurulum hatası
        raise RuntimeError(
            "pdf okumak için 'pypdf' gerekli: pip install -r requirements.txt"
        ) from exc

    reader = PdfReader(str(path))
    sayfalar: list[str] = []
    for sayfa in reader.pages:
        metin = (sayfa.extract_text() or "").strip()
        if metin:
            sayfalar.append(_pdf_satirlarini_birlestir(metin))
    return "\n\n".join(sayfalar)


def _pdf_satirlarini_birlestir(metin: str) -> str:
    """PDF'in sayfa genişliğine göre kırdığı satırları paragrafa geri çevir.

    PDF'te her görsel satır ayrı bir satır sonu taşır; bunlar olduğu gibi
    bırakılırsa her satır ayrı bir "paragraf" sanılır ve chunk'lama cümleleri
    ortadan böler. Tek satır sonları boşluğa çevrilir, boş satırlar paragraf
    sınırı olarak korunur.
    """
    paragraflar = _PARAGRAPH_SPLIT.split(metin)
    duzeltilmis = []
    for paragraf in paragraflar:
        satirlar = [s.strip() for s in paragraf.splitlines() if s.strip()]
        if satirlar:
            duzeltilmis.append(" ".join(satirlar))
    return "\n\n".join(duzeltilmis)


def iter_corpus_files(corpus_dir: Path | None = None) -> list[Path]:
    """Corpus klasöründeki desteklenen dosyalar, kararlı sırada."""
    root = Path(corpus_dir) if corpus_dir else config.CORPUS_DIR
    if not root.exists():
        return []
    return sorted(
        p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
    )


def _split_long(text: str, max_chars: int) -> list[str]:
    """Üst sınırı aşan tek parçayı cümlelere, gerekirse karaktere böl."""
    pieces: list[str] = []
    buffer = ""
    for sentence in _SENTENCE_SPLIT.split(text):
        candidate = f"{buffer} {sentence}".strip() if buffer else sentence
        if len(candidate) <= max_chars:
            buffer = candidate
            continue
        if buffer:
            pieces.append(buffer)
        # Tek cümle bile sığmıyorsa sert kes.
        while len(sentence) > max_chars:
            pieces.append(sentence[:max_chars])
            sentence = sentence[max_chars:]
        buffer = sentence
    if buffer:
        pieces.append(buffer)
    return pieces


def chunk_text(
    text: str,
    max_chars: int | None = None,
    overlap_chars: int | None = None,
) -> list[str]:
    """Metni chunk listesine çevir. Boş girdi boş liste döner."""
    max_chars = max_chars if max_chars is not None else config.CHUNK_MAX_CHARS
    overlap = overlap_chars if overlap_chars is not None else config.CHUNK_OVERLAP_CHARS
    if overlap >= max_chars:
        raise ValueError("overlap_chars, max_chars'tan küçük olmalı")

    paragraphs = [p.strip() for p in _PARAGRAPH_SPLIT.split(text) if p.strip()]
    if not paragraphs:
        return []

    units: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) <= max_chars:
            units.append(paragraph)
        else:
            units.extend(_split_long(paragraph, max_chars))

    chunks: list[str] = []
    current = ""
    for unit in units:
        candidate = f"{current}\n\n{unit}" if current else unit
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            chunks.append(current)
            tail = current[-overlap:] if overlap else ""
            current = f"{tail}\n\n{unit}".strip() if tail else unit
            # Devreden kuyruk yüzünden sınır aşıldıysa kuyruğu at.
            if len(current) > max_chars:
                current = unit
        else:
            current = unit
    if current:
        chunks.append(current)
    return chunks


def kullanilabilir(chunk: str) -> bool:
    """Chunk anlamlı metin mi, yoksa çıkarım artığı mı?

    YALNIZCA okunamaz glif kodları elenir. Sayı yoğun parçalar (tablolar,
    ölçüm sonuçları) BİLEREK elenmez — içlerinde gerçek cevap olabilir.
    """
    return len(_GLIF_KODU.findall(chunk)) < _GLIF_ESIGI


def chunk_file(path: Path, corpus_dir: Path | None = None) -> tuple[str, list[str]]:
    """Dosyayı oku ve (kaynak adı, chunk'lar) döndür.

    Kaynak adı corpus köküne göre göreli yoldur — cevaplarda kullanıcıya bu gösterilir.
    """
    root = Path(corpus_dir) if corpus_dir else config.CORPUS_DIR
    try:
        source = str(path.relative_to(root)).replace("\\", "/")
    except ValueError:
        source = path.name
    return source, chunk_text(read_text(path))
